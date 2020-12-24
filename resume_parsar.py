import datetime
import math
import os
import re
import sys
from pathlib import Path

import PyPDF2
import docx
import en_core_web_sm
import numpy as np
import pandas as pd
import textract
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from word2number import w2n
#from resume_parser import resumeparse
#import tika
#tika.initVM()
#from tika import parser

nlp = en_core_web_sm.load()

from nltk.corpus import stopwords
STOPWORDS = set(stopwords.words('english'))
curdir = os.getcwd()
sys.path.append(curdir)
class resumeParsar():
    def __init__(self):
        self.full_path = os.getcwd()
        self.data_path = str(Path(self.full_path).parents[0]) + '/app/data/'
        self.df = pd.read_csv(self.data_path+'Skills.csv')
        self.skills = list(self.df['skills'])
        #self.skills = [' {0} '.format(elem) for elem in self.skills]
        self.df_hobbies = pd.read_csv(self.data_path+ 'Hobbies.csv')
        self.hobbies = list(self.df_hobbies['hobbies'])
        self.hobbies = [' {0} '.format(elem) for elem in self.hobbies]
        self.education = ['B.E.', 'B.E', 'B.S',
                 'M.E', 'M.E.', 'M.S',
                'BTECH', 'B.TECH','B.TECH', 'M.TECH',
                'MTECH','BCA','MCA','BSC','MSC','B.S.C','M.S.C','M.C.A' 
                ' SSC ', ' HSC ', 'CBSE', 'ICSE ', ' Xth ', ' XIIth ','10TH','12TH ',
                'BACHELORS OF ENGINEERING','B.E',' PG ',' PGP ','PGPA','BBA','12th' ,'10th ',' CA '
                ,' CBSE ',' ICSE '
               ,' PGDBA ',' UNIVERISTY ',' SCHOOL ',' COLLEGE ']
        self.df_locations = pd.read_csv(self.data_path + 'Locations.csv')
        self.locations = list(self.df_locations['Cities'])
        self.locations = [' {0} '.format(elem) for elem in self.locations]


    def listToString(self,ls):
        str1 = ""
        for ele in ls:
            str1 += ele.strip() + ','
        return str1

    def getFileExtension(self,filename):
        extension = filename.split('.')[-1]
        return extension

    def checkExtension(self,filename):
        ValidExtension = ["pdf", "docx","doc"]
        extension = self.getFileExtension(filename)
        if extension not in ValidExtension:
            print("Error:Not a valid File")
            return False
        else:
            return True

    def readpdfFile(self,filename):
        content = ""
        pdfFileObj = open(filename, 'rb')
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
        noofpages = pdfReader.numPages
        for i in range(0, noofpages):
            page = pdfReader.getPage(i)
            content_page = page.extractText()
            content = content + ' ' + content_page
        content = content.encode('utf-8')
        return content

    def iter_hyperlink_rels(self,rels):
        hls = ""
        for rel in rels:
            if rels[rel].reltype == RT.HYPERLINK:
                hls = hls + str(rels[rel]._target)
        return hls

    def readtables(self,filename):
        document = docx.Document(filename)
        tables = document.tables
        data = []
        if len(tables) > 0:
            data = []
            for table in tables:
                keys = None
                for i, row in enumerate(table.rows):
                    text = (cell.text for cell in row.cells)

                    if i == 0:
                        keys = tuple(text)
                        continue
                    row_data = dict(zip(keys, text))
                    data.append(row_data)
        return data

    def extracttabletextskills(self,filename):
        document = docx.Document(filename)
        tables = document.tables
        data = []
        table_text = ""
        if len(tables) > 0:
            data = []
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        table_text = table_text + " " + paragraph.text
            if any(skill in str(table_text).lower() for skill in self.skills):
                break
            return table_text

    def extracttabletexteducation(self,filename):
        document = docx.Document(filename)
        tables = document.tables
        data = []
        table_text = ""
        if len(tables) > 0:
            data = []
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        table_text = table_text + " " + paragraph.text
            if any(edu in str(table_text).lower() for edu in self.education):
                break
            return table_text
        else:
            return ''

    def readdocxFile(self,filename):
        doc = docx.Document(filename)
        fullText = ""

        for para in doc.paragraphs:
            text_para = (para.text)
            fullText = fullText + ' ' + text_para
        rels = doc.part.rels
        hls = self.iter_hyperlink_rels(rels)
        fullText = fullText + " " + hls
        return fullText

    def doc_to_docx(self,filename):
        text = textract.process(filename)
        text = text.decode("utf-8")
        texts = text.split('\n')
        doc = docx.Document()
        para_change_flag = 0
        para_text = ''

        for text in texts:
            # print(text)
            if text.strip() == '':
                para_change_flag = 1
            else:
                para_change_flag = 0
            if para_change_flag == 1 and para_text.strip() != '':
                doc_para = doc.add_paragraph(para_text)
                para_text = ''
            else:
                para_text = para_text + '\n' + text
        new_filename = filename + 'x'
        doc.save(new_filename)
        return new_filename

    def extract_name(self,fulltext):
        name = re.findall("[\dA-Za-z+\' ']*", fulltext)[0]
        return name

    def extract_mobile_number(self,text):
        phone = re.findall(re.compile('([0-9]{10}|[0-9]{4}\s[0-9]{3}\s[0-9]{3})'), text)
        if phone:
            number = ''.join(phone[0])
            if len(number) > 10:
                return '+' + number
            else:
                return number

    def extract_email(self,text):
        email = re.findall(r"[a-z0-9\.\-+_]+@[a-z0-9\.\-+_]+\.[a-z]+", text)
        if email:
            try:
                return email
            except IndexError:
                return None

    def extract_primary_secondry_skill(self,filename):
        primary_skill = []
        secondry_skill = []
        count = 0
        doc = docx.Document(filename)
        for para in doc.paragraphs:
            text_para = (para.text)
            if any(skill in str(text_para).lower() for skill in self.skills) and count < 2:
                primary_skill.extend([skill for skill in self.skills if (skill in str(text_para).lower())])
                count = count + 1
            elif any(skill in str(text_para).lower() for skill in self.skills) and count > 2:
                secondry_skill.extend([skill for skill in self.skills if (skill in str(text_para).lower())])
                count = count + 1
            if len(primary_skill) > 3:
                primary_skill_final = primary_skill[0:2]
                secondry_skill.extend(primary_skill[3:len(primary_skill) - 1])
            else:
                primary_skill_final=primary_skill
        primary_skill = list(dict.fromkeys(primary_skill_final))
        secondry_skill = list(dict.fromkeys(secondry_skill))
        #return str(primary_skill).replace('[','').replace(']',''), str(secondry_skill).replace('[','').replace(']','')
        return self.listToString(primary_skill),self.listToString(secondry_skill)

    def extract_current_preferred_location(self,filename):
        count = 0
        Current_Location=[]
        Preferred_Location=[]
        doc = docx.Document(filename)
        for para in doc.paragraphs:
            text_para = (para.text)
            if any(str(location) in str(text_para).lower() for location in self.locations) and count < 2:
                Current_Location = [str(location) for location in self.locations if (str(location) in str(text_para).lower())]
                count = count + 1
            elif any(str(location) in str(text_para).lower() for location in self.locations) and count >= 2:
                Preferred_Location = [str(location) for location in self.locations if (str(location) in str(text_para).lower())]
                count = count + 1
        Current_Location = list(set(Current_Location))
        Preferred_Location = list(set(Preferred_Location))
        return self.listToString(Current_Location), self.listToString(Preferred_Location)

    def extract_location(self,resume_text):
        location_all = []
        if any(str(location) in str(resume_text).lower() for location in self.locations):
            location_all = [str(location) for location in self.locations if (str(location) in str(resume_text).lower())]
        return self.listToString(location_all)

    def extract_hobbies(self,resume_text):
        hobbies_all = []
        if any(str(hobby) in str(resume_text).lower() for hobby in self.hobbies):
            hobbies_all = [str(hobby) for hobby in self.hobbies if (str(hobby) in str(resume_text).lower())]
        return self.listToString(hobbies_all)

    def visa_check(self,resume_text):
        visa_all=[]
        visas = ['h1b','h1n1','l1','schengen']
        if any(str(visa) in str(resume_text).lower() for visa in visas):
            visa_all = [str(visa) for visa in visas if (str(visa) in str(resume_text).lower())]
        return self.listToString(visa_all)

    def extract_skills(self,resume_text):
        skills_all = ''
        if any(skill in str(resume_text).lower() for skill in self.skills):
            skills_all = [skill for skill in self.skills if (skill in str(resume_text).lower())]
        return self.listToString(skills_all)

    def extract_education(self,filename):
        ls_edu = []
        doc = docx.Document(filename)
        escape_text = ['personal', 'course', 'certification', 'certifications', 'certificate',
                       'declaration', 'declare',
                       'skill', 'project', 'experience', 'projects','company','companies']
        escape_text.extend(self.skills)
        for para in doc.paragraphs:
            text_para = (para.text)
            if any(ext in text_para.lower() for ext in escape_text):
                ls_edu = ls_edu
            elif any(edu.lower() in text_para.lower() for edu in self.education):
                ls_edu.append(text_para.strip())
        if len(ls_edu) > 0:
            return self.listToString(ls_edu)
        else:
            return self.listToString(ls_edu)

    def getexpr(self,filename):
        text_exp = ""
        exclude_text=['name','father','address','personal','married','single']
        exclude_text.extend(self.education)
        exclude_text.extend(self.locations)
        doc = docx.Document(filename)
        for para in doc.paragraphs:
            text_para = (para.text)
            match = re.findall(r'.*(\s[1-3][0-9]{3}|3000)', text_para)
            if match:
                if any(ext.lower() in text_para.lower() for ext in exclude_text):
                    text_exp = text_exp
                else:
                    text_exp = text_exp + "\n" + text_para + " "
        return text_exp

    def project_details(self,filename):
        doc = docx.Document(filename)
        project_text = ''
        exclude_texts = ['education', 'skill', 'certificate', 'windows', 'win', 'version', 'university', 'school',
                         'college']
        exclude_texts.extend(self.education)
        start_para = 6
        para_size = len(doc.paragraphs)
        end_para = math.ceil(para_size * 80 / 100)
        for para in doc.paragraphs[start_para:end_para]:
            text_para = (para.text)
            if any(et.lower() in str(text_para).lower() for et in exclude_texts):
                project_text = project_text
            else:
                project_text = project_text + ' ' + text_para
        return project_text

    def get_project_regex(self,filename):
        pro_text = []
        project_texts = ['project']
        exclude_texts = ['education', 'skill', 'experience']
        exclude_texts.extend(self.education)
        doc = docx.Document(filename)
        for para in doc.paragraphs:
            text_para = (para.text)
            if any(pt in str(text_para).lower() for pt in project_texts):
                if any(et in str(text_para).lower() for et in exclude_texts):
                    pro_text = pro_text
                else:
                    pro_text.append(text_para.replace('project', ''))
        return pro_text

    def extractoverall_experience_through_yrs(self,filename):
        doc = docx.Document(filename)
        now = datetime.datetime.now()
        curr_yr = now.year
        max_working_yr = curr_yr - 40
        yrs = []
        exclude_texts = ['education', 'skill', 'certificate', 'windows', 'win', 'version']
        exclude_texts.extend(self.education)
        for para in doc.paragraphs:
            text_para = (para.text)
            match = re.findall(r'.*(\s[1-3][0-9]{3}|3000)', text_para)
            if match:
                if any(et.lower() in str(text_para).lower() for et in exclude_texts):
                    yrs = yrs
                else:
                    yrs.extend(match)
        if len(yrs) > 0:
            yrs = [str(yr).strip() for yr in yrs if
                   (str(yr).strip() > str(max_working_yr) and str(yr).strip() <= str(curr_yr))]
            yrs.append(curr_yr)
            yrs = list(set(yrs))
            yrs = [int(yr) for yr in yrs]
            yrs.sort()
            ex_limt = np.mean(np.diff(yrs)) + 2
            ex_limit_pos = []
            pos = 0
            for i in np.diff(yrs):
                if i > ex_limt:
                    ex_limit_pos.append(pos)
                pos = pos + 1
            for pos in ex_limit_pos:
                del yrs[pos]
            overall_exp = max(yrs) - min(yrs)
        else:
            overall_exp = ''
        return str(overall_exp) + ' years'

    def extract_overall_experience(self,filename):
        doc = docx.Document(filename)
        yrs=[]
        for para in doc.paragraphs:
            text_para = (para.text)
            exp_tokens = ['yrs','years','year','exp','experience']
            for para in doc.paragraphs:
                text_para = (para.text)
                yrs=[int(i) for i in text_para.split() if i.isdigit() and int(i)<40]
                try:
                    flag = w2n.word_to_num(text_para)
                    if flag > 100:
                        flag = 0
                except:
                    flag = 0
                if any(ext in text_para.lower() for ext in exp_tokens) and (len(yrs)>0 or flag >0):
                    yrs.append(flag)
        if len(yrs)>0 and max(yrs) >0:
            return str(max(yrs))+' '+'years'
        else:
            return self.extractoverall_experience_through_yrs(filename)
            #return ''

    def get_file_nm(self, filenm):
        try:
            filenms = filenm.split('/')
            length = len(filenms) - 1
            file_name = filenms[length]
        except:
            file_name = filenm
        return file_name

    def parse_direct(self,filename):
        '''try:
            result_direct = resumeparse.read_file(filename)
            return result_direct
        except:'''
        return {"degree":[],"designition":[],"email":"","name":"","phone":"","skills":[],"total_exp":'NA',"university":[]}

    def generate_resume_result(self,filename):
        result = {}
        if self.checkExtension(filename):
            extension = self.getFileExtension(filename)
            if extension == 'docx' or extension == 'doc':
                if extension == 'doc':
                    filename = self.doc_to_docx(filename)
                fulltext = self.readdocxFile(filename)
                tables = self.readtables(filename)
            else:
                fulltext = self.readpdfFile(filename)
            tabular_data = self.readtables(filename)
            result_direct = self.parse_direct(filename)
            if self.extract_name(fulltext) and self.extract_name(fulltext).strip()!='':
                result['Name'] = self.extract_name(fulltext)
            else:
                result['Name'] = str(result_direct['name']).replace('Phone','').replace('phone','').replace('Mobile','').replace('mobile','')

            if self.extract_mobile_number(fulltext) and self.extract_mobile_number(fulltext).strip()!='':
                result['Mobile_number'] = self.extract_mobile_number(fulltext)
            else:
                result['Mobile_number'] = result_direct['phone']
            if self.extract_email(fulltext) and len(self.extract_email(fulltext)) > 0:
                result['Email'] = self.extract_email(fulltext)
            else:
                result['Email'] = result_direct['email']
            if self.extract_skills(fulltext):
                result['Skills_All'] = self.extract_skills(fulltext)
            elif len(self.extracttabletextskills(fulltext, self.skills)) > 0:
                result['Skills_All'] = self.extracttabletextskills(fulltext, self.skills)
            if self.extract_education(filename) and len(self.extract_education(filename)) >0:
                result['Education'] = self.extract_education(filename)
            elif len(self.extracttabletexteducation(filename)) > 0:
                result['Education'] = self.extracttabletexteducation(filename)
            else:
                result['Education'] = result_direct['degree']
            result['Experience'] = self.getexpr(filename)
            Primary_skills, Secondary_skills = self.extract_primary_secondry_skill(filename)
            Current_Location,Preferred_Location = self.extract_current_preferred_location(filename)
            Overall_Experience = self.extract_overall_experience(filename)
            result['Primary_Skills'] = Primary_skills
            result['Secondary_Skills'] = Secondary_skills
            result['Current_Location'] = Current_Location
            result['Preferred_Location'] = self.extract_location(fulltext)
            result['Projects'] = self.project_details(filename)
            if Overall_Experience and Overall_Experience.strip()!='':
                result['Overall_Experience'] = Overall_Experience
            else:
                result['Overall_Experience'] = result_direct['total_exp']
            result['Designation'] = result_direct['designition']
            result['Hobbies'] = self.extract_hobbies(fulltext)
            result['Visa'] = self.visa_check(fulltext)
            result['Filename'] = self.get_file_nm(filename)
        return result




if __name__ == "__main__":
    rp = resumeParsar()
    # filename= "//home/lid/Downloads//Omar_Nour_CV.docx"
    filename = sys.argv[1]
    print(filename)
    result = rp.generate_resume_result(filename)
    print(result)

