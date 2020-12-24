import PyPDF2
import textract
import docx
import pandas as pd
from nltk.stem import WordNetLemmatizer
from nltk.corpus import stopwords
import re
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import spacy
from spacy.matcher import Matcher
import spacy
import glob
from pathlib import Path
import os
from gensim import corpora, models, similarities
from nltk.corpus import stopwords
import sys
STOPWORDS = set(stopwords.words('english'))

class jdParsar():
    def __init__(self):
        self.full_path = os.getcwd()
        self.data_path = str(Path(self.full_path).parents[0]) + '/app/data/'
        self.df = pd.read_csv(self.data_path + 'Skills.csv')
        self.skills = list(self.df['skills'])
        self.df_locations = pd.read_csv(self.data_path + 'Locations.csv')
        self.locations = list(self.df_locations['Cities'])
        self.education = [
            'BE', 'B.E.', 'B.E', 'BS', 'B.S',
            'ME', 'M.E', 'M.E.', 'MS', 'M.S',
            'BTECH', 'B.TECH', 'M.TECH', 'MTECH', 'BCA', 'MCA', 'BSC', 'MSC', 'B.S.C', 'M.S.C', 'M.C.A'
                                                                                                'SSC', 'HSC', 'CBSE',
            'ICSE', 'X', 'XII', '10TH', '12TH', 'BACHELORS OF ENGINEERING', 'B.E', 'PG', 'PGP', 'PGPA'
            , 'PGDBA']
        self.designations=['software engineer','trainee','intern','manager','data analyst','data engineer'
                          'architect','test engineer','qa','data scientist','analyst','consultant',
                          'business analyst','director','president','vice president']
        self.job_types=['contractor','contract','parttime','part time','fulltime','full time',
                        'third party payroll','remotly','wfh']

    def extract_jd_title(self,filename):
        title = ''
        doc = docx.Document(filename)
        fullText = ""
        for para in doc.paragraphs:
            text_para = (para.text)
            if len(text_para) > 5:
                title = text_para
                break
        return title

    def extract_job_type(self,resume_text):
        job_types_all = ''
        if any(job_type in str(resume_text).lower() for job_type in self.job_types):
            job_types_all = job_types_all + ',' + str(
                [job_type for job_type in self.job_types if (job_type in str(resume_text).lower())]).replace('[', '').replace(']',
                                                                                                             '').replace(
                '\'', '').replace(', ', ',').strip()
        return job_types_all

    def extract_skills(self,resume_text):
        skills_all = ''
        if any(skill in str(resume_text).lower() for skill in self.skills):
            skills_all = [skill for skill in self.skills if (skill in str(resume_text).lower())]
        return self.listToString(skills_all)

    def extract_designation(self,resume_text):
        designations_all = ''
        if any(designation in str(resume_text).lower() for designation in self.designations):
            designations_all = [designation for designation in self.designations if (designation in str(resume_text).lower())]
        return self.listToString(designations_all)

    def extract_location(self,resume_text):
        location_all = ''
        if any(str(location) in str(resume_text).lower() for location in self.locations):
            location_all = [str(location) for location in self.locations if (str(location) in str(resume_text).lower())]
        return self.listToString(location_all)
    def extract_notice_period(self,resumetext):
        if 'immediate' in resumetext.lower() and 'join' in resumetext.lower():
            np='immediate'
        elif 'day' in resumetext.lower() and len(re.findall(r"\d{2}", resumetext.lower()))>0:
            np=re.findall(r"\d{2}", resumetext.lower())
            np=str(np)+'days'
        elif 'month' in resumetext.lower() and len(re.findall(r"\d{1}", resumetext.lower()))>0:
            np=re.findall(r"\d{1}", resumetext.lower())
            np=str(np)+'months'
        elif 'month' in resumetext.lower() and len(re.findall(r"\d{1}", resumetext.lower()))>0:
            np=re.findall(r"\d{1}", resumetext.lower())
            np=str(np)+'days'
            return np

    def listToString(self,ls):
        str1 = ""
        for ele in ls:
            str1 += ele.strip() + ','
        return str1

    def extract_primary_secondry_skill(self, filename):
        primary_skill = []
        secondry_skill = []
        count = 0
        doc = docx.Document(filename)
        for para in doc.paragraphs:
            text_para = (para.text)
            if any(skill in str(text_para).lower() for skill in self.skills) and count < 2:
                primary_skill.extend([skill for skill in self.skills if (skill in str(text_para).lower())])
                count = count + 1
            elif any(skill in str(text_para).lower() for skill in self.skills) and count > 2:
                secondry_skill.extend([skill for skill in self.skills if (skill in str(text_para).lower())])
                count = count + 1
            if len(primary_skill) > 3:
                primary_skill_final = primary_skill[0:2]
                secondry_skill.extend(primary_skill[3:len(primary_skill) - 1])
            else:
                primary_skill_final = primary_skill
        primary_skill = list(dict.fromkeys(primary_skill_final))
        secondry_skill = list(dict.fromkeys(secondry_skill))
        #return str(primary_skill).replace('[', '').replace(']', ''), str(secondry_skill).replace('[', '').replace(']','')
        return self.listToString(primary_skill),self.listToString(secondry_skill)

    def iter_hyperlink_rels(self,rels):
        hls = ""
        for rel in rels:
            if rels[rel].reltype == RT.HYPERLINK:
                hls = hls + str(rels[rel]._target)
        return hls

    def readtables(self,filename):
        document = docx.Document(filename)
        tables = document.tables
        data = []
        if len(tables) > 0:
            data = []
            for table in tables:
                keys = None
                for i, row in enumerate(table.rows):
                    text = (cell.text for cell in row.cells)

                    if i == 0:
                        keys = tuple(text)
                        continue
                    row_data = dict(zip(keys, text))
                    data.append(row_data)
        return data

    def extracttabletextskills(self,filename, skills):
        document = docx.Document(filename)
        tables = document.tables
        data = []
        table_text = ""
        if len(tables) > 0:
            data = []
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        table_text = table_text + " " + paragraph.text
            if any(skill in str(table_text).lower() for skill in skills):
                break
            return table_text

    def extracttabletexteducation(self,filename):
        document = docx.Document(filename)
        tables = document.tables
        data = []
        table_text = ""
        if len(tables) > 0:
            data = []
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        table_text = table_text + " " + paragraph.text
            if any(edu in str(table_text).lower() for edu in self.education):
                break
            return table_text

    def readdocxFile(self,filename):
        doc = docx.Document(filename)
        fullText = ""
        for para in doc.paragraphs:
            text_para = (para.text)
            fullText = fullText + ' ' + text_para
        rels = doc.part.rels
        hls = self.iter_hyperlink_rels(rels)
        fullText = fullText + " " + hls
        return fullText
    def getfilenames(self,jd_location):
        file_names = glob.glob(jd_location)
        return file_names

    def extract_skills_with_experience(self,filename):
        swe = ''
        doc = docx.Document(filename)
        for para in doc.paragraphs:
            text_para = (para.text)
            search_string = ['yr', 'yrs', 'year', 'years', 'exp', 'experience']
            if any(search in str(text_para).lower() for search in search_string):
                # swe=swe+','+str([search for search in search_string if(search in str(text_para).lower())]).replace('[','').replace(']','').replace('\'','').replace(', ',',').strip()
                swe = swe + " " + text_para
        return swe
    def getparsedjd_multiple(self,file_names):
        fulltexts = []
        titles = []
        skills_all = []
        swe = []
        primary_skills = []
        secondry_skills = []
        for file in file_names:
            resume_text = self.readdocxFile(file)
            fulltexts.append(resume_text)
            titles.append(self.extract_jd_title(file))
            skills_all.append(self.extract_skills(resume_text))
            swe.append(self.extract_skills_with_experience(file))
            primary_skill, secondry_skill = self.extract_primary_secondry_skill(file)
            primary_skills.append(primary_skill)
            secondry_skills.append(secondry_skill)
        df_jds = pd.DataFrame()
        df_jds['title'] = titles
        df_jds['skills_all'] = skills_all
        df_jds['swe'] = swe
        df_jds['primary_skills'] = primary_skills
        df_jds['secondry_skills'] = secondry_skills
        return df_jds
    def get_file_nm(self, filenm):
        try:
            filenms = filenm.split('/')
            length = len(filenms) - 1
            file_name = filenms[length]
        except:
            file_name = filenm
        return file_name

    def getparsedjd(self,file):
        result={}
        resume_text = self.readdocxFile(file)
        file_name = self.get_file_nm(file)
        result['Filename'] = file_name
        result['Title'] = self.extract_jd_title(file)
        result['Skills_All'] = self.extract_skills(resume_text)
        primary_skill, secondry_skill = self.extract_primary_secondry_skill(file)
        result['Primary_Skills'] = primary_skill
        result['Secondry_Skills'] = secondry_skill
        result['Designation'] = self.extract_designation(resume_text)
        result['Skill_With_Experince'] = self.extract_skills_with_experience(file)
        result['Job_Type']=self.extract_job_type(resume_text)
        result['Location']=self.extract_location(resume_text)
        result['Notice_Period']=self.extract_notice_period(resume_text)
        return result

if __name__ == "__main__":
    jdp = jdParsar()
    filename = sys.argv[1]
    # filename = "/home/lid/Downloads/Job Description-20201121T112409Z-001/Job Description/Django Engineer_JD.docx"
    print(jdp.getparsedjd(filename))




